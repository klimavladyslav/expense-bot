import os
import json
import httpx
import tempfile
import subprocess
from datetime import datetime
from zoneinfo import ZoneInfo

KYIV_TZ = ZoneInfo("Europe/Kyiv")


def extract_text_from_pdf(path: str) -> str:
    import pdfplumber
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
            # Also try to extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    text += " | ".join(str(c) for c in row if c) + "\n"
    return text


def extract_text_from_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def extract_text_from_doc(path: str) -> str:
    """Convert .doc to .docx via LibreOffice, then parse."""
    tmp_dir = tempfile.mkdtemp()
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", tmp_dir, path],
            capture_output=True, timeout=30
        )
        base = os.path.splitext(os.path.basename(path))[0]
        docx_path = os.path.join(tmp_dir, base + ".docx")
        if os.path.exists(docx_path):
            return extract_text_from_docx(docx_path)
        # Fallback: try antiword or just read raw
        return result.stdout.decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[DOC parse error: {e}]"


def extract_text_from_excel(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            if any(c is not None for c in row):
                lines.append(" | ".join(str(c) if c is not None else "" for c in row))
    return "\n".join(lines)


def extract_text(path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".doc":
        return extract_text_from_doc(path)
    elif ext in (".xlsx", ".xls", ".xlsm"):
        return extract_text_from_excel(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


PARSE_SYSTEM_PROMPT = """You are an invoice parser. Extract structured data from invoice/packing list documents.

Return ONLY valid JSON, no markdown, no explanation:

{
  "invoices": [
    {
      "invoice_no": "string",
      "invoice_date": "YYYY-MM-DD or null",
      "supplier": "string",
      "currency": "USD|EUR|UAH",
      "invoice_amount_total": number,
      "products": [
        {
          "product": "product description",
          "qty": number,
          "unit": "pcs|sets|kg|other",
          "unit_price": number,
          "cbm": number or null
        }
      ]
    }
  ],
  "multiple_containers": false,
  "note": "any important note or null"
}

Rules:
- If multiple invoices clearly relate to ONE container/shipment, put them all under one invoice entry (use the main invoice number)
- If they are clearly separate containers, return multiple invoice objects and set multiple_containers: true
- If unclear whether same or different container, set multiple_containers: null (will ask user)
- Extract CBM per product line if available in packing list, otherwise null
- invoice_amount_total is the grand total of the invoice
- Dates in YYYY-MM-DD format

IMPORTANT - Grouped/bundled product lines:
- Some invoices list components of a set across multiple rows but with ONE shared quantity and ONE shared price for the whole block (e.g. a bedding set listed as: row1=duvet cover, row2=flat sheet, row3=pillow case — all sharing qty=4080 and unit_price=$11.10)
- In this case, treat the entire block as ONE product line, not multiple
- Product name format: "[block name] ([component1], [component2], [component3])"
  Example: "4 pcs set double size (duvet cover, flat sheet, 2x pillow case)"
- Use the block-level qty and unit_price (the ones that appear once for the whole group)
- Each distinct block with its own qty/price = one product line
- Only apply this grouping when multiple rows genuinely share a single qty+price; if each row has its own qty and price, keep them separate
"""


async def parse_invoice_file(path: str, filename: str) -> dict:
    text = extract_text(path, filename)

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    payload = {
        "model": "claude-sonnet-4-5",
        "max_tokens": 2000,
        "system": PARSE_SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": f"Parse this invoice document:\n\n{text[:8000]}"}]
    }

    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01"
            },
            json=payload
        )
        response.raise_for_status()
        data = response.json()

    raw = data["content"][0]["text"].strip()
    if "```" in raw:
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)
